"""
Generate Clinician Notes in Word (.docx) format for RareSense.AI Demo.
Each note is crafted to trigger NLP entity extraction (symptoms, medications, diagnoses, lab values)
matching the HPO codes in the NLP engine.
"""
import sys
import os
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

# ─────────────────────────────────────────────────────────────────
# Clinical Notes — designed to trigger NLP extraction
# ─────────────────────────────────────────────────────────────────
CLINICAL_NOTES = [
    {
        "title": "Case 1 — Systemic Lupus Erythematosus (SLE) Presentation",
        "patient": "Elena Vasquez, 28F",
        "specialty": "Rheumatology",
        "physician": "Dr. Sarah Mitchell",
        "note_type": "Consultation Note",
        "date": "January 10, 2025",
        "text": (
            "28F presents with erythematous malar rash across both cheeks and reports photosensitivity for the past 3 months. "
            "Rash worsens with sun exposure. Patient also complains of chronic fatigue, joint pain in bilateral hands and wrists, "
            "and oral ulcers on the buccal mucosa. Physical examination reveals classic butterfly rash sparing nasolabial folds. "
            "ANA titer 1:320. Anti-dsDNA antibodies elevated. ESR 52 mm/hr. CRP 28.4 mg/L. "
            "Platelet count 98,000/uL (low — thrombocytopenia noted). Hemoglobin 10.2 g/dL (mild anemia). "
            "Assessment: findings consistent with Systemic Lupus Erythematosus. "
            "Started hydroxychloroquine 200mg BID and prednisone 10mg daily. "
            "Advised strict sun protection. Follow-up in 4 weeks with repeat labs."
        ),
    },
    {
        "title": "Case 2 — SLE Follow-Up with Renal Involvement",
        "patient": "Elena Vasquez, 28F",
        "specialty": "Nephrology",
        "physician": "Dr. James Chen",
        "note_type": "Progress Note",
        "date": "March 22, 2025",
        "text": (
            "Follow-up visit for known lupus patient. Patient reports worsening fatigue and new onset peripheral edema "
            "in bilateral lower extremities. Urinalysis reveals proteinuria (2+). eGFR 62 mL/min (reduced). "
            "Creatinine 1.6 mg/dL (elevated). 24-hour urine collection shows 2.8g protein — consistent with nephritis. "
            "Complement levels low (C3 and C4 depressed). Platelet count 82,000/uL. "
            "Joint pain persists despite hydroxychloroquine. Reports occasional fever and night sweats. "
            "Assessment: lupus nephritis Class III suspected. Plan: renal biopsy scheduled. "
            "Increased prednisone to 40mg daily with taper plan. Added mycophenolate 500mg BID. "
            "Continue hydroxychloroquine. Monitor renal function closely."
        ),
    },
    {
        "title": "Case 3 — Cystic Fibrosis Infant Presentation",
        "patient": "Liam Miller, 8-month-old Male",
        "specialty": "Pulmonology",
        "physician": "Dr. James Carter",
        "note_type": "Consultation Note",
        "date": "February 15, 2025",
        "text": (
            "8-month-old male infant referred for evaluation of chronic cough and failure to thrive. "
            "Parents report persistent thick respiratory mucus and recurrent lower respiratory tract infections. "
            "Mother notes 'salty-tasting skin' when kissing the infant. Weight below 5th percentile despite adequate caloric intake. "
            "Sweat chloride test: 82 mmol/L (diagnostic threshold >60). "
            "CFTR genotyping: homozygous F508del mutation confirmed. "
            "Chest X-ray shows early bronchiectatic changes. Pancreatic insufficiency suspected given steatorrhea and weight loss. "
            "Assessment: Cystic Fibrosis confirmed. "
            "Plan: Initiate dornase alfa for airway clearance. Start pancreatic enzyme replacement therapy (PERT). "
            "Chest physiotherapy twice daily. Nutritional supplementation with fat-soluble vitamins. "
            "Pulmonary function monitoring quarterly."
        ),
    },
    {
        "title": "Case 4 — Huntington Disease Neurological Evaluation",
        "patient": "Robert Chen, 44M",
        "specialty": "Neurology",
        "physician": "Dr. Allison House",
        "note_type": "Consultation Note",
        "date": "April 5, 2025",
        "text": (
            "44M presents with progressive involuntary choreic movements of upper and lower extremities over the past 18 months. "
            "Patient reports significant cognitive decline, difficulty with executive function, and increasing memory loss. "
            "Family reports personality changes and mood disturbances including irritability and depression. "
            "Father deceased at age 52 from 'neurological disease' — family history strongly positive for autosomal dominant inheritance. "
            "Neurological examination reveals diffuse chorea, motor impersistence, and mild ataxia. "
            "Cognitive assessment: MMSE 22/30 showing decline. "
            "Genetic testing: 42 CAG repeats in the HTT gene (pathogenic; normal <36). "
            "MRI brain shows caudate atrophy with mild cortical thinning. "
            "Assessment: Huntington Disease confirmed. Muscle weakness noted in proximal upper extremities. "
            "Patient also reports dysphagia with solid foods and unintentional weight loss of 12 lbs over 6 months. "
            "Plan: Initiate tetrabenazine for chorea management. Speech therapy for dysphagia assessment. "
            "Genetic counseling for family members. Neuropsychological follow-up in 3 months."
        ),
    },
    {
        "title": "Case 5 — Systemic Sclerosis (Scleroderma) Workup",
        "patient": "Maria Santos, 52F",
        "specialty": "Rheumatology",
        "physician": "Dr. Ana Rodriguez",
        "note_type": "Consultation Note",
        "date": "May 18, 2025",
        "text": (
            "52F referred for evaluation of progressive skin tightening and Raynaud phenomenon. "
            "Patient reports cold-triggered digital color changes (white → blue → red) for 2 years, worsening fatigue, "
            "and shortness of breath on exertion. Physical examination reveals sclerodactyly, telangiectasia on face and hands, "
            "and digital pitting scars. Bilateral joint pain in hands and wrists noted. "
            "Pulmonary function tests show restrictive pattern with reduced DLCO — suggesting pulmonary fibrosis. "
            "High-resolution CT chest: early ground-glass opacities in bilateral lower lobes. "
            "Echocardiogram: estimated RVSP 42 mmHg — mild pulmonary hypertension. "
            "Reports progressive dysphagia with solid foods over 6 months. "
            "ANA positive (1:640, nucleolar pattern). Anti-Scl-70 (anti-topoisomerase I) positive. "
            "ESR 38 mm/hr. CRP 18.5 mg/L. "
            "Assessment: Systemic Sclerosis (diffuse cutaneous variant). "
            "Plan: Start mycophenolate 500mg BID for skin and lung disease. Continue amlodipine for Raynaud. "
            "Proton pump inhibitor for esophageal dysmotility. Annual echocardiogram for pulmonary hypertension screening."
        ),
    },
    {
        "title": "Case 6 — Gaucher Disease Evaluation",
        "patient": "David Petrov, 35M",
        "specialty": "Hematology",
        "physician": "Dr. Michael Brown",
        "note_type": "Consultation Note",
        "date": "June 12, 2025",
        "text": (
            "35M referred for unexplained splenomegaly and persistent thrombocytopenia. "
            "Patient reports chronic fatigue for 2+ years, easy bruising, and occasional arthralgia in hips and knees. "
            "Physical examination confirms palpable splenomegaly (6cm below costal margin) and hepatomegaly. "
            "CBC: Platelet count 68,000/uL. Hemoglobin 10.8 g/dL (mild anemia). WBC count 3.8k (low). "
            "Peripheral smear: pancytopenia confirmed with no blasts. "
            "Acid phosphatase elevated. Chitotriosidase markedly elevated (diagnostic biomarker). "
            "Bone marrow biopsy: characteristic Gaucher cells (lipid-laden macrophages) identified. "
            "Beta-glucocerebrosidase enzyme activity: 2.1 nmol/hr/mg (severely reduced; normal >8.0). "
            "GBA1 gene sequencing: compound heterozygous N370S/L444P mutations. "
            "Assessment: Gaucher Disease Type 1 confirmed. "
            "Plan: Initiate enzyme replacement therapy (ERT) with imiglucerase infusions every 2 weeks. "
            "Baseline bone density scan. Monitor CBC, liver function, and spleen size quarterly."
        ),
    },
    {
        "title": "Case 7 — Behçet Disease with Multi-System Involvement",
        "patient": "Tariq Al-Hassan, 31M",
        "specialty": "Rheumatology",
        "physician": "Dr. David Kim",
        "note_type": "Consultation Note",
        "date": "July 8, 2025",
        "text": (
            "31M presenting with recurrent oral ulcers (≥3 episodes in past 12 months) and bilateral painful genital ulcerations. "
            "Patient reports severe headache, photosensitivity, and arthralgia in bilateral knees and ankles. "
            "Ophthalmological examination reveals bilateral anterior uveitis with cells and flare. "
            "Skin examination shows erythema nodosum-like lesions on lower extremities — consistent with skin rash. "
            "Pathergy test positive (papule >2mm at 48 hours). "
            "ESR 44 mm/hr. CRP 32.6 mg/L. HLA-B51 positive. "
            "Reports chronic fatigue and episodes of low-grade fever. "
            "Evidence of vasculitis on skin biopsy — neutrophilic infiltration of vessel walls. "
            "Assessment: Behçet Disease with ocular, mucocutaneous, and articular involvement. "
            "Plan: Colchicine 0.5mg BID for oral ulcers and arthritis. Topical corticosteroids for genital ulcers. "
            "Azathioprine 50mg daily for uveitis prevention. Ophthalmology follow-up in 2 weeks."
        ),
    },
    {
        "title": "Case 8 — Fabry Disease Screening",
        "patient": "Wei Nakamura, 28M",
        "specialty": "Neurology",
        "physician": "Dr. Thomas Green",
        "note_type": "Consultation Note",
        "date": "August 20, 2025",
        "text": (
            "28M presents with burning paresthesia and numbness in hands and feet since adolescence. "
            "Patient reports episodes of severe abdominal pain, chronic fatigue, and intermittent fever during flares. "
            "Family history: maternal uncle with renal failure at age 40, mother with 'unexplained neuropathy'. "
            "Physical examination reveals small dark-red angiokeratomas on trunk and bathing trunk distribution. "
            "Reduced sweating reported (hypohidrosis). Progressive headache and cognitive difficulties noted. "
            "Creatinine 1.4 mg/dL. eGFR 72 (mildly reduced). Proteinuria (1+) on urinalysis. "
            "Echocardiogram: concentric left ventricular hypertrophy — early cardiomyopathy. "
            "Alpha-galactosidase A enzyme activity: 0.8 nmol/hr/mg (markedly reduced; normal >3.0). "
            "GLA gene sequencing: hemizygous p.R227Q missense mutation identified. "
            "Assessment: Fabry Disease confirmed. "
            "Plan: Initiate enzyme replacement therapy with agalsidase beta. ACE inhibitor (lisinopril) for renal protection. "
            "Cardiology referral for surveillance. Screen family members."
        ),
    },
    {
        "title": "Case 9 — Mixed Connective Tissue Disease",
        "patient": "Priya Sharma, 38F",
        "specialty": "Rheumatology",
        "physician": "Dr. Emily Watson",
        "note_type": "Consultation Note",
        "date": "September 3, 2025",
        "text": (
            "38F presents with overlapping features of multiple autoimmune conditions. "
            "Chief complaints: severe arthralgia and joint swelling in bilateral MCPs, PIPs, and wrists. "
            "Reports Raynaud phenomenon with clear triphasic digital color changes. "
            "Progressive fatigue and myalgia with proximal muscle weakness in upper extremities over 6 months. "
            "Physical exam: swollen sausage-like digits, periungual erythema, and mild malar rash. "
            "Pulmonary: mild shortness of breath on exertion. CXR shows subtle interstitial markings. "
            "Labs: ANA 1:1280 (speckled pattern). Anti-U1 RNP strongly positive (>200 units). "
            "Anti-dsDNA negative. Anti-Smith negative. CK mildly elevated at 380 U/L. "
            "ESR 48 mm/hr. CRP 22.8 mg/L. "
            "Assessment: Mixed Connective Tissue Disease (Sharp Syndrome). "
            "Plan: Start prednisone 20mg daily taper for acute inflammation. "
            "Methotrexate 15mg weekly as steroid-sparing agent. "
            "Amlodipine 5mg for Raynaud management. Pulmonary function test baseline."
        ),
    },
    {
        "title": "Case 10 — Sickle Cell Disease Crisis",
        "patient": "Aisha Williams, 22F",
        "specialty": "Hematology",
        "physician": "Dr. Lisa Park",
        "note_type": "Progress Note",
        "date": "October 14, 2025",
        "text": (
            "22F with known sickle cell disease (HbSS) admitted for acute vaso-occlusive crisis. "
            "Presenting with severe bilateral joint pain in hips, knees, and lower back. "
            "Reports fatigue, chest pain, and shortness of breath. Temperature 38.7°C — fever present. "
            "Physical exam: mild jaundice noted in sclera. Palpable splenomegaly. "
            "CBC: Hemoglobin 6.8 g/dL (severe anemia). WBC count 14.2k (elevated). Platelet count 420,000/uL. "
            "Reticulocyte count 12.4% (elevated, hemolytic state). LDH elevated at 580 U/L. "
            "Chest X-ray: no acute infiltrate — no evidence of acute chest syndrome currently. "
            "Peripheral smear confirms sickle cells and target cells. "
            "Assessment: Sickle Cell Disease vaso-occlusive crisis with acute hemolytic anemia. "
            "Plan: IV fluid hydration. Morphine PCA for pain management. "
            "Incentive spirometry to prevent acute chest syndrome. Monitor for splenic sequestration. "
            "Consider hydroxyurea initiation after acute episode resolves. Folic acid 1mg daily."
        ),
    },
]


def create_docx(output_path: str):
    """Create a professional Word document with clinician notes."""
    doc = Document()

    # ── Styles ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1D, 0x1D, 0x1F)

    # ── Title Page ──
    doc.add_paragraph('')
    doc.add_paragraph('')
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('RareSense.AI')
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Clinical Case Notes for NLP Demonstration')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run(
        'LLM-Powered Rare Disease Detection from Unstructured Clinical Notes\n'
        'MongoDB + BioBERT + HPO Phenotype Matching'
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_paragraph('')
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_paragraph('')
    usage = doc.add_paragraph()
    usage.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = usage.add_run(
        'USAGE: Copy-paste any note below into the RareSense.AI "Add Clinical Note" form.\n'
        'The NLP engine will automatically extract symptoms, medications, diagnoses, and lab values.\n'
        'Then run the Matching Engine to find rare disease matches.'
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
    run.italic = True

    doc.add_page_break()

    # ── Table of Contents ──
    toc_heading = doc.add_heading('Table of Contents', level=1)
    for i, note in enumerate(CLINICAL_NOTES, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}.  {note["title"]}')
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        meta = p.add_run(f'    [{note["specialty"]} — {note["physician"]}]')
        meta.font.size = Pt(9)
        meta.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_page_break()

    # ── Clinical Notes ──
    for i, note in enumerate(CLINICAL_NOTES, 1):
        # Heading
        doc.add_heading(f'{note["title"]}', level=1)

        # Metadata table
        meta_table = doc.add_table(rows=5, cols=2)
        meta_table.style = 'Light List Accent 1'
        
        fields = [
            ('Patient', note['patient']),
            ('Specialty', note['specialty']),
            ('Physician', note['physician']),
            ('Note Type', note['note_type']),
            ('Date', note['date']),
        ]
        for row_idx, (label, value) in enumerate(fields):
            cell_label = meta_table.rows[row_idx].cells[0]
            cell_value = meta_table.rows[row_idx].cells[1]
            
            run_label = cell_label.paragraphs[0].add_run(label)
            run_label.bold = True
            run_label.font.size = Pt(10)
            
            run_value = cell_value.paragraphs[0].add_run(value)
            run_value.font.size = Pt(10)

        doc.add_paragraph('')

        # Note label
        label_p = doc.add_paragraph()
        run = label_p.add_run('CLINICAL NOTE TEXT (copy this into the system):')
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

        # Note text in a bordered paragraph
        note_p = doc.add_paragraph()
        note_p.style = doc.styles['Normal']
        run = note_p.add_run(note['text'])
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1D, 0x1D, 0x1F)
        
        # Formatting for the note paragraph
        fmt = note_p.paragraph_format
        fmt.left_indent = Cm(0.5)
        fmt.right_indent = Cm(0.5)
        fmt.space_before = Pt(6)
        fmt.space_after = Pt(12)

        # Separator
        if i < len(CLINICAL_NOTES):
            sep = doc.add_paragraph()
            sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = sep.add_run('─' * 60)
            run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
            run.font.size = Pt(8)
            doc.add_paragraph('')

    # ── Quick Reference: Expected Extractions ──
    doc.add_page_break()
    doc.add_heading('Quick Reference: Expected NLP Extractions', level=1)

    ref_note = doc.add_paragraph()
    run = ref_note.add_run(
        'Below are the key entities the RareSense NLP engine should extract from each note. '
        'Use this to validate the system is working correctly during the demo.'
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    run.italic = True

    expected = [
        ("Case 1 — SLE", "Symptoms: Malar rash, Photosensitivity, Chronic fatigue, Arthralgia, Oral ulcers, Thrombocytopenia, Anemia\nMedications: Hydroxychloroquine, Prednisone\nDiagnoses: Systemic Lupus Erythematosus\nLab Values: ANA Titer 320, ESR 52, CRP 28.4, Platelet Count 98000, Hemoglobin 10.2"),
        ("Case 2 — Nephritis", "Symptoms: Chronic fatigue, Peripheral edema, Proteinuria, Nephritis, Arthralgia, Fever, Night sweats, Thrombocytopenia\nMedications: Hydroxychloroquine, Prednisone, Mycophenolate\nDiagnoses: Lupus\nLab Values: eGFR 62, Creatinine 1.6, Platelet Count 82000"),
        ("Case 3 — CF", "Symptoms: Cough, Weight loss\nMedications: None detected\nDiagnoses: Cystic Fibrosis\nLab Values: None"),
        ("Case 4 — Huntington", "Symptoms: Cognitive decline, Ataxia, Muscle weakness, Dysphagia, Weight loss, Tremor\nMedications: None detected\nDiagnoses: Huntington\nLab Values: None"),
        ("Case 5 — Scleroderma", "Symptoms: Raynaud, Chronic fatigue, Dyspnea, Telangiectasia, Arthralgia, Pulmonary fibrosis, Pulmonary hypertension, Dysphagia, Skin rash\nMedications: Mycophenolate, Amlodipine\nDiagnoses: Systemic Sclerosis / Scleroderma\nLab Values: ESR 38, CRP 18.5"),
        ("Case 6 — Gaucher", "Symptoms: Splenomegaly, Thrombocytopenia, Hepatomegaly, Chronic fatigue, Arthralgia, Anemia, Pancytopenia\nMedications: None\nDiagnoses: Gaucher\nLab Values: Platelet Count 68000, Hemoglobin 10.8, WBC Count 3.8"),
        ("Case 7 — Behçet", "Symptoms: Oral ulcers, Headache, Photosensitivity, Arthralgia, Uveitis, Skin rash, Chronic fatigue, Fever, Vasculitis\nMedications: Azathioprine\nDiagnoses: Behçet\nLab Values: ESR 44, CRP 32.6"),
        ("Case 8 — Fabry", "Symptoms: Paresthesia, Numbness, Abdominal pain, Chronic fatigue, Fever, Headache, Cardiomyopathy, Proteinuria\nMedications: Lisinopril\nDiagnoses: Fabry\nLab Values: Creatinine 1.4, eGFR 72"),
        ("Case 9 — MCTD", "Symptoms: Arthralgia, Joint swelling, Raynaud, Chronic fatigue, Myalgia, Muscle weakness, Malar rash, Dyspnea\nMedications: Prednisone, Methotrexate, Amlodipine\nDiagnoses: Mixed Connective Tissue\nLab Values: ESR 48, CRP 22.8"),
        ("Case 10 — SCD", "Symptoms: Arthralgia, Chronic fatigue, Chest pain, Dyspnea, Fever, Jaundice, Splenomegaly, Anemia\nMedications: Folic acid\nDiagnoses: Sickle Cell\nLab Values: Hemoglobin 6.8, WBC Count 14.2, Platelet Count 420000"),
    ]

    ref_table = doc.add_table(rows=len(expected) + 1, cols=2)
    ref_table.style = 'Light List Accent 1'
    
    # Header row
    ref_table.rows[0].cells[0].paragraphs[0].add_run('Case').bold = True
    ref_table.rows[0].cells[1].paragraphs[0].add_run('Expected Extractions').bold = True

    for row_idx, (case, extractions) in enumerate(expected, 1):
        ref_table.rows[row_idx].cells[0].paragraphs[0].add_run(case).font.size = Pt(9)
        run = ref_table.rows[row_idx].cells[1].paragraphs[0].add_run(extractions)
        run.font.size = Pt(8)

    # Save
    doc.save(output_path)
    print(f"Document saved to: {output_path}")


if __name__ == '__main__':
    output = os.path.join(
        os.path.dirname(os.path.dirname(__file__)),
        'RareSense_Clinical_Notes.docx'
    )
    create_docx(output)
